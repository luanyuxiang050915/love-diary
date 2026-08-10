# -*- coding: utf-8 -*-
"""把 测试用例.md 转换为 恋爱日记测试用例.docx（compact reference guide 风格）。

用法：python build_testcases_docx.py
依赖：python-docx（Codex 自带 Python 环境已包含）
"""
import re

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

PINK = RGBColor(0xF8, 0xA5, 0xC2)
HOT = RGBColor(0xFF, 0x6B, 0x9D)
DARK = RGBColor(0x33, 0x33, 0x33)
GRAY = RGBColor(0x88, 0x88, 0x88)
FONT = "微软雅黑"

doc = Document()

# 页面：A4，留白适中
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(1.8)
sec.bottom_margin = Cm(1.8)
sec.left_margin = Cm(1.8)
sec.right_margin = Cm(1.8)

normal = doc.styles["Normal"]
normal.font.name = FONT
normal.font.size = Pt(10.5)
normal.font.color.rgb = DARK
normal.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
normal.paragraph_format.space_after = Pt(4)
normal.paragraph_format.line_spacing = 1.3


def set_run(r, size=10.5, bold=False, color=DARK, italic=False):
    r.font.name = FONT
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    r.font.color.rgb = color
    r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)


def add_rich(p, text, size=10.5, bold=False, color=DARK):
    for part in re.split(r"(\*\*[^*]+\*\*)", text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            set_run(p.add_run(part[2:-2]), size, True, color)
        else:
            set_run(p.add_run(part), size, bold, color)
    return p


def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run(text), 22, True, HOT)
    p.paragraph_format.space_after = Pt(2)


def add_h1(text):
    p = doc.add_paragraph()
    set_run(p.add_run(text), 15, True, HOT)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), "F8A5C2")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_h2(text):
    p = doc.add_paragraph()
    set_run(p.add_run(text), 12, True, PINK)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def add_table(rows):
    ncols = len(rows[0])
    if ncols == 4:
        widths = [Cm(3.8), Cm(5.2), Cm(5.2), Cm(2.0)]
    else:
        widths = [Cm(14.2 / ncols)] * ncols
    t = doc.add_table(rows=len(rows), cols=ncols)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ri, row in enumerate(rows):
        for ci, cell_text in enumerate(row):
            cell = t.cell(ri, ci)
            cell.width = widths[ci]
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            if ri == 0:
                shade(cell, "FDEAF0")
                set_run(p.add_run(cell_text), 9.5, True, HOT)
            else:
                set_run(p.add_run(cell_text), 9.5, False, DARK)
            if ci == ncols - 1:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def main():
    with open("测试用例.md", encoding="utf-8") as f:
        lines = f.read().splitlines()

    i = 0
    pending = []
    while i < len(lines):
        line = lines[i]
        if line.startswith("|"):
            pending.append(line)
        else:
            if pending:
                rows = []
                for l in pending:
                    cells = [c.strip() for c in l.strip().strip("|").split("|")]
                    if all(set(c) <= {"-", ":"} for c in cells):
                        continue
                    rows.append(cells)
                add_table(rows)
                pending = []

            s = line.strip()
            if not s or s == "---":
                i += 1
                continue
            if s.startswith("# "):
                add_title(s[2:])
            elif s.startswith("## "):
                add_h1(s[3:])
            elif s.startswith("### "):
                add_h2(s[4:])
            elif s.startswith("> "):
                p = doc.add_paragraph()
                add_rich(p, s[2:], 10, False, GRAY)
                p.paragraph_format.space_after = Pt(6)
            elif s.startswith("- [ ] "):
                p = doc.add_paragraph(style="List Bullet")
                set_run(p.add_run("[ ] "), 10.5, True, PINK)
                add_rich(p, s[6:], 10.5)
            elif re.match(r"^\d+\.\s", s):
                p = doc.add_paragraph(style="List Number")
                add_rich(p, re.sub(r"^\d+\.\s", "", s), 10.5)
            else:
                add_rich(doc.add_paragraph(), s, 10.5)
        i += 1

    if pending:
        rows = []
        for l in pending:
            cells = [c.strip() for c in l.strip().strip("|").split("|")]
            if all(set(c) <= {"-", ":"} for c in cells):
                continue
            rows.append(cells)
        add_table(rows)

    doc.save("恋爱日记测试用例.docx")
    print("DOCX saved: 恋爱日记测试用例.docx")


if __name__ == "__main__":
    main()
